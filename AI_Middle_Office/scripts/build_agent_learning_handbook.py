from __future__ import annotations

import argparse
import html
import re
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PAGE_WIDTH_IN = 8.5
PAGE_HEIGHT_IN = 11.0
MARGIN_IN = 1.0
HEADER_FOOTER_IN = 0.492
USABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120

ASCII_FONT = "Calibri"
EAST_ASIA_FONT = "Microsoft YaHei"
CODE_FONT = "Consolas"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
TEXT = "243447"
MUTED = "667085"
LIGHT_BLUE = "E8EEF5"
PALE_BLUE = "F4F7FB"
CODE_BG = "F6F8FA"
BORDER = "D0D7DE"
WHITE = "FFFFFF"


@dataclass(frozen=True)
class Note:
    title: str
    source: Path
    relative_source: str
    chapter_index: int
    chapter_title: str
    subcategory: str | None
    bookmark: str


@dataclass
class Chapter:
    index: int
    title: str
    bookmark: str
    notes: list[Note]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    if not widths_dxa:
        return
    delta = USABLE_WIDTH_DXA - sum(widths_dxa)
    widths_dxa[-1] += delta

    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:w"), str(USABLE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        row.height = None
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[min(idx, len(widths_dxa) - 1)])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if table.rows:
        tr_pr = table.rows[0]._tr.get_or_add_trPr()
        table_header = tr_pr.find(qn("w:tblHeader"))
        if table_header is None:
            table_header = OxmlElement("w:tblHeader")
            tr_pr.append(table_header)
        table_header.set(qn("w:val"), "true")


def set_paragraph_border(paragraph, side: str, color: str, size: int = 8, space: int = 4) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    border = p_bdr.find(qn(f"w:{side}"))
    if border is None:
        border = OxmlElement(f"w:{side}")
        p_bdr.append(border)
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), str(size))
    border.set(qn("w:space"), str(space))
    border.set(qn("w:color"), color)


def set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_run_font(run, size: float | None = None, bold: bool | None = None, color: str | None = None,
                 ascii_font: str = ASCII_FONT, east_asia_font: str = EAST_ASIA_FONT,
                 italic: bool | None = None) -> None:
    run.font.name = ascii_font
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia_font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def add_page_field(paragraph, field_name: str) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = field_name
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)
    set_run_font(run, size=8.5, color=MUTED)


def add_bookmark(paragraph, name: str, bookmark_id: int) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def add_internal_hyperlink(paragraph, text: str, anchor: str, color: str = BLUE) -> None:
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    hyperlink.set(qn("w:history"), "1")
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), ASCII_FONT)
    r_fonts.set(qn("w:hAnsi"), ASCII_FONT)
    r_fonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    color_node = OxmlElement("w:color")
    color_node.set(qn("w:val"), color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(r_fonts)
    r_pr.append(color_node)
    r_pr.append(underline)
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_external_hyperlink(paragraph, text: str, url: str) -> None:
    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    hyperlink.set(qn("w:history"), "1")
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), ASCII_FONT)
    r_fonts.set(qn("w:hAnsi"), ASCII_FONT)
    r_fonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    color_node = OxmlElement("w:color")
    color_node.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(r_fonts)
    r_pr.append(color_node)
    r_pr.append(underline)
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


INLINE_PATTERN = re.compile(
    r"(\*\*.+?\*\*|__.+?__|~~.+?~~|`[^`]+`|\[[^\]]+\]\([^)]+\)|https?://[^\s<>()]+)"
)


def add_inline_content(paragraph, text: str, current_source: Path | None,
                       bookmark_by_source: dict[Path, str]) -> None:
    text = html.unescape(text)
    position = 0
    for match in INLINE_PATTERN.finditer(text):
        if match.start() > position:
            set_run_font(paragraph.add_run(text[position:match.start()]))
        token = match.group(0)
        if token.startswith(("**", "__")):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, bold=True)
        elif token.startswith("~~"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run)
            run.font.strike = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=9.5, ascii_font=CODE_FONT, east_asia_font=EAST_ASIA_FONT)
            run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
        elif token.startswith("["):
            match_link = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token)
            if not match_link:
                set_run_font(paragraph.add_run(token))
            else:
                label, url = match_link.groups()
                if url.startswith(("http://", "https://")):
                    add_external_hyperlink(paragraph, label, url)
                elif current_source is not None and not url.startswith("#"):
                    target_url = url.split("#", 1)[0]
                    target = (current_source.parent / target_url).resolve()
                    anchor = bookmark_by_source.get(target)
                    if anchor:
                        add_internal_hyperlink(paragraph, label, anchor)
                    else:
                        set_run_font(paragraph.add_run(label), color=BLUE)
                else:
                    set_run_font(paragraph.add_run(label), color=BLUE)
        else:
            url = token.rstrip(".,;，。；")
            suffix = token[len(url):]
            add_external_hyperlink(paragraph, url, url)
            if suffix:
                set_run_font(paragraph.add_run(suffix))
        position = match.end()
    if position < len(text):
        set_run_font(paragraph.add_run(text[position:]))


def strip_frontmatter(text: str) -> str:
    lines = text.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    if lines and lines[0].strip() == "---":
        for idx in range(1, len(lines)):
            if lines[idx].strip() == "---":
                return "\n".join(lines[idx + 1:])
    return "\n".join(lines)


def split_table_row(line: str) -> list[str]:
    row = line.strip()
    if row.startswith("|"):
        row = row[1:]
    if row.endswith("|"):
        row = row[:-1]
    cells = re.split(r"(?<!\\)\|", row)
    return [cell.strip().replace(r"\|", "|") for cell in cells]


def is_table_separator(line: str) -> bool:
    cells = split_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells)


def compute_column_widths(rows: list[list[str]]) -> list[int]:
    column_count = max(len(row) for row in rows)
    weights: list[float] = []
    for col in range(column_count):
        lengths = []
        for row in rows[:40]:
            value = row[col] if col < len(row) else ""
            lengths.append(max(2, min(40, len(re.sub(r"[*_`]", "", value)))))
        weights.append(max(5.0, sum(lengths) / max(1, len(lengths))))
    min_width = 600 if column_count >= 8 else 780
    available = USABLE_WIDTH_DXA - min_width * column_count
    if available <= 0:
        base = USABLE_WIDTH_DXA // column_count
        return [base for _ in range(column_count)]
    weight_sum = sum(weights)
    return [min_width + int(available * weight / weight_sum) for weight in weights]


def add_markdown_table(document: Document, rows: list[list[str]], current_source: Path,
                       bookmark_by_source: dict[Path, str]) -> None:
    if not rows:
        return
    column_count = max(len(row) for row in rows)
    normalized = [row + [""] * (column_count - len(row)) for row in rows]
    table = document.add_table(rows=len(normalized), cols=column_count)
    table.style = "Table Grid"
    widths = compute_column_widths(normalized)
    set_table_geometry(table, widths)
    font_size = 8.5 if column_count <= 5 else 7.5 if column_count <= 8 else 6.5
    for row_index, row in enumerate(normalized):
        for column_index, value in enumerate(row):
            cell = table.cell(row_index, column_index)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.05
            add_inline_content(paragraph, value, current_source, bookmark_by_source)
            for run in paragraph.runs:
                set_run_font(run, size=font_size, bold=row_index == 0)
            if row_index == 0:
                set_cell_shading(cell, LIGHT_BLUE)
            elif row_index % 2 == 0:
                set_cell_shading(cell, "FAFBFC")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_numbering_definition(document: Document, ordered: bool) -> int:
    numbering = document.part.numbering_part.element
    cache = getattr(document, "_handbook_numbering_cache", None)
    if cache is None:
        abstract_ids = [
            int(element.get(qn("w:abstractNumId")))
            for element in numbering.findall(qn("w:abstractNum"))
            if element.get(qn("w:abstractNumId")) is not None
        ]
        num_ids = [
            int(element.get(qn("w:numId")))
            for element in numbering.findall(qn("w:num"))
            if element.get(qn("w:numId")) is not None
        ]
        cache = {
            "next_abstract_id": max(abstract_ids, default=0) + 1,
            "next_num_id": max(num_ids, default=0) + 1,
            "abstracts": {},
            "shared_bullet_num_id": None,
        }
        setattr(document, "_handbook_numbering_cache", cache)

    key = "ordered" if ordered else "bullet"
    abstract_id = cache["abstracts"].get(key)
    if abstract_id is None:
        abstract_id = cache["next_abstract_id"]
        cache["next_abstract_id"] += 1
        cache["abstracts"][key] = abstract_id

        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "hybridMultilevel")
        abstract.append(multi)
        for level in range(9):
            lvl = OxmlElement("w:lvl")
            lvl.set(qn("w:ilvl"), str(level))
            start = OxmlElement("w:start")
            start.set(qn("w:val"), "1")
            num_fmt = OxmlElement("w:numFmt")
            num_fmt.set(qn("w:val"), "decimal" if ordered else "bullet")
            lvl_text = OxmlElement("w:lvlText")
            if ordered:
                lvl_text.set(qn("w:val"), f"%{level + 1}.")
            else:
                lvl_text.set(qn("w:val"), ("•", "◦", "▪")[level % 3])
            suffix = OxmlElement("w:suff")
            suffix.set(qn("w:val"), "space")
            p_pr = OxmlElement("w:pPr")
            tabs = OxmlElement("w:tabs")
            tab = OxmlElement("w:tab")
            tab.set(qn("w:val"), "num")
            tab.set(qn("w:pos"), str(540 + level * 360))
            tabs.append(tab)
            indentation = OxmlElement("w:ind")
            indentation.set(qn("w:left"), str(540 + level * 360))
            indentation.set(qn("w:hanging"), "270")
            p_pr.append(tabs)
            p_pr.append(indentation)
            lvl.append(start)
            lvl.append(num_fmt)
            lvl.append(lvl_text)
            lvl.append(suffix)
            lvl.append(p_pr)
            abstract.append(lvl)
        numbering.append(abstract)

    if not ordered and cache["shared_bullet_num_id"] is not None:
        return cache["shared_bullet_num_id"]

    num_id = cache["next_num_id"]
    cache["next_num_id"] += 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    if ordered:
        for level in range(9):
            override = OxmlElement("w:lvlOverride")
            override.set(qn("w:ilvl"), str(level))
            start_override = OxmlElement("w:startOverride")
            start_override.set(qn("w:val"), "1")
            override.append(start_override)
            num.append(override)
    numbering.append(num)
    if not ordered:
        cache["shared_bullet_num_id"] = num_id
    return num_id


def apply_numbering(paragraph, num_id: int, level: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(min(level, 8)))
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_node)


def add_code_block(document: Document, code: str, language: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.08)
    paragraph.paragraph_format.right_indent = Inches(0.02)
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.05
    set_paragraph_shading(paragraph, CODE_BG)
    set_paragraph_border(paragraph, "left", BLUE, size=10, space=5)
    if language:
        label = paragraph.add_run(language.upper() + "\n")
        set_run_font(label, size=7.5, bold=True, color=MUTED, ascii_font=CODE_FONT)
    code_run = paragraph.add_run(code.rstrip())
    set_run_font(code_run, size=8.0, color=TEXT, ascii_font=CODE_FONT, east_asia_font=EAST_ASIA_FONT)


def add_quote(document: Document, lines: list[str], current_source: Path,
              bookmark_by_source: dict[Path, str]) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.right_indent = Inches(0.05)
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(7)
    paragraph.paragraph_format.line_spacing = 1.2
    set_paragraph_shading(paragraph, PALE_BLUE)
    set_paragraph_border(paragraph, "left", BLUE, size=12, space=6)
    for idx, line in enumerate(lines):
        if idx:
            paragraph.add_run("\n")
        add_inline_content(paragraph, line, current_source, bookmark_by_source)
    for run in paragraph.runs:
        set_run_font(run, size=10.0, color=DARK_BLUE)


def add_horizontal_rule(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(5)
    set_paragraph_border(paragraph, "bottom", BORDER, size=4, space=1)


class MarkdownRenderer:
    def __init__(self, document: Document, bookmark_by_source: dict[Path, str]):
        self.document = document
        self.bookmark_by_source = bookmark_by_source
        self.list_kind: str | None = None
        self.list_num_id: int | None = None

    def end_list(self) -> None:
        self.list_kind = None
        self.list_num_id = None

    def render(self, text: str, source: Path, note_heading_level: int) -> None:
        lines = strip_frontmatter(text).split("\n")
        index = 0
        paragraph_buffer: list[str] = []
        first_h1_skipped = False

        def flush_paragraph() -> None:
            nonlocal paragraph_buffer
            if not paragraph_buffer:
                return
            value = " ".join(part.strip() for part in paragraph_buffer if part.strip())
            paragraph_buffer = []
            if not value:
                return
            paragraph = self.document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing = 1.25
            add_inline_content(paragraph, value, source, self.bookmark_by_source)
            self.end_list()

        while index < len(lines):
            line = lines[index]
            stripped = line.strip()

            if stripped.startswith("```"):
                flush_paragraph()
                language = stripped[3:].strip()
                index += 1
                code_lines = []
                while index < len(lines) and not lines[index].strip().startswith("```"):
                    code_lines.append(lines[index])
                    index += 1
                add_code_block(self.document, "\n".join(code_lines), language)
                self.end_list()
                index += 1
                continue

            heading_match = re.match(r"^(#{1,6})\s+(.+)$", stripped)
            if heading_match:
                flush_paragraph()
                source_level = len(heading_match.group(1))
                title = heading_match.group(2).strip().rstrip("#").strip()
                if source_level == 1 and not first_h1_skipped:
                    first_h1_skipped = True
                    index += 1
                    continue
                target_level = min(6, note_heading_level + max(1, source_level - 1))
                paragraph = self.document.add_paragraph(style=f"Heading {target_level}")
                add_inline_content(paragraph, title, source, self.bookmark_by_source)
                self.end_list()
                index += 1
                continue

            if stripped.startswith(">"):
                flush_paragraph()
                quote_lines = []
                while index < len(lines) and lines[index].strip().startswith(">"):
                    quote_lines.append(re.sub(r"^\s*>\s?", "", lines[index]))
                    index += 1
                add_quote(self.document, quote_lines, source, self.bookmark_by_source)
                self.end_list()
                continue

            if (
                "|" in line
                and index + 1 < len(lines)
                and "|" in lines[index + 1]
                and is_table_separator(lines[index + 1])
            ):
                flush_paragraph()
                header = split_table_row(line)
                index += 2
                rows = [header]
                while index < len(lines) and "|" in lines[index] and lines[index].strip():
                    rows.append(split_table_row(lines[index]))
                    index += 1
                add_markdown_table(self.document, rows, source, self.bookmark_by_source)
                self.end_list()
                continue

            list_match = re.match(r"^(\s*)([-*+]|\d+[.)])\s+(.+)$", line)
            if list_match:
                flush_paragraph()
                indent, marker, content = list_match.groups()
                ordered = marker[0].isdigit()
                kind = "ordered" if ordered else "bullet"
                if kind != self.list_kind:
                    self.list_kind = kind
                    self.list_num_id = add_numbering_definition(self.document, ordered)
                paragraph = self.document.add_paragraph()
                paragraph.paragraph_format.space_after = Pt(4)
                paragraph.paragraph_format.line_spacing = 1.25
                level = min(8, max(0, len(indent.replace("\t", "    ")) // 2))
                apply_numbering(paragraph, self.list_num_id or 1, level)
                content = re.sub(r"^\[ \]\s*", "☐ ", content)
                content = re.sub(r"^\[[xX]\]\s*", "☑ ", content)
                add_inline_content(paragraph, content, source, self.bookmark_by_source)
                index += 1
                continue

            if stripped in {"---", "***", "___"}:
                flush_paragraph()
                add_horizontal_rule(self.document)
                self.end_list()
                index += 1
                continue

            if not stripped:
                flush_paragraph()
                self.end_list()
                index += 1
                continue

            paragraph_buffer.append(line)
            index += 1

        flush_paragraph()


def configure_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = ASCII_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.widow_control = True

    heading_specs = {
        1: (16, BLUE, 18, 10),
        2: (13, BLUE, 14, 7),
        3: (12, DARK_BLUE, 10, 5),
        4: (11, "365F91", 8, 4),
        5: (10.5, "4F5F6F", 7, 3),
        6: (10, "5F6B78", 6, 3),
    }
    for level, (size, color, before, after) in heading_specs.items():
        style = styles[f"Heading {level}"]
        style.font.name = ASCII_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    for style_name in ("Header", "Footer"):
        style = styles[style_name]
        style.font.name = ASCII_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
        style.font.size = Pt(8.5)
        style.font.color.rgb = RGBColor.from_string(MUTED)


def configure_page(document: Document) -> None:
    for section in document.sections:
        section.page_width = Inches(PAGE_WIDTH_IN)
        section.page_height = Inches(PAGE_HEIGHT_IN)
        section.top_margin = Inches(MARGIN_IN)
        section.bottom_margin = Inches(MARGIN_IN)
        section.left_margin = Inches(MARGIN_IN)
        section.right_margin = Inches(MARGIN_IN)
        section.header_distance = Inches(HEADER_FOOTER_IN)
        section.footer_distance = Inches(HEADER_FOOTER_IN)
        section.different_first_page_header_footer = True


def configure_headers_and_footers(document: Document) -> None:
    for section in document.sections:
        header = section.header
        paragraph = header.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_paragraph_border(paragraph, "bottom", LIGHT_BLUE, size=4, space=3)
        run = paragraph.add_run("AI AGENT ENGINEER / FDE · STUDY HANDBOOK")
        set_run_font(run, size=8, bold=True, color=MUTED)

        footer = section.footer
        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_border(paragraph, "top", LIGHT_BLUE, size=4, space=3)
        run = paragraph.add_run("第 ")
        set_run_font(run, size=8.5, color=MUTED)
        add_page_field(paragraph, "PAGE")
        run = paragraph.add_run(" 页 / 共 ")
        set_run_font(run, size=8.5, color=MUTED)
        add_page_field(paragraph, "NUMPAGES")
        run = paragraph.add_run(" 页")
        set_run_font(run, size=8.5, color=MUTED)


def parse_readme(notes_root: Path) -> list[Chapter]:
    readme = (notes_root / "README.md").read_text(encoding="utf-8")
    chapters: list[Chapter] = []
    current_chapter: Chapter | None = None
    current_subcategory: str | None = None
    note_index = 0

    for line in readme.splitlines():
        chapter_match = re.match(r"^###\s+(\d+)\.\s+(.+?)\s*$", line)
        if chapter_match:
            index = int(chapter_match.group(1))
            title = chapter_match.group(2).strip()
            current_chapter = Chapter(index=index, title=title, bookmark=f"chapter_{index:02d}", notes=[])
            chapters.append(current_chapter)
            current_subcategory = None
            continue
        subcategory_match = re.match(r"^####\s+(.+?)\s*$", line)
        if subcategory_match and current_chapter is not None:
            current_subcategory = subcategory_match.group(1).strip()
            continue
        note_match = re.match(r"^-\s+\[([^\]]+)\]\(([^)]+\.md)\)\s*$", line)
        if note_match and current_chapter is not None:
            note_index += 1
            title, relative = note_match.groups()
            relative = relative.removeprefix("./")
            source = (notes_root / relative).resolve()
            if not source.exists():
                raise FileNotFoundError(f"README 中的笔记不存在：{relative}")
            note = Note(
                title=title.strip(),
                source=source,
                relative_source=relative.replace("\\", "/"),
                chapter_index=current_chapter.index,
                chapter_title=current_chapter.title,
                subcategory=current_subcategory,
                bookmark=f"note_{note_index:03d}",
            )
            current_chapter.notes.append(note)
    if len(chapters) != 14:
        raise ValueError(f"预期 14 个章节，实际 {len(chapters)}")
    total_notes = sum(len(chapter.notes) for chapter in chapters)
    if total_notes != 77:
        raise ValueError(f"预期 77 篇笔记，实际 {total_notes}")
    return chapters


CHAPTER_DESCRIPTIONS = {
    1: "先建立岗位目标、知识地图与复习优先级。",
    2: "理解生产级 Agent 的分层架构与完整数据流。",
    3: "掌握状态、可靠性、上下文和任务恢复等运行时基础。",
    4: "补齐模型推理、Prompt、解码和微调的必要原理。",
    5: "把工具调用从“能调”升级为“安全、可控、可验证”。",
    6: "理解决策循环、规划、工作流和执行边界。",
    7: "设计短期、工作、长期记忆及其生命周期。",
    8: "构建可评测、可追溯、有证据链的生产级 RAG。",
    9: "补齐 MySQL、Redis、消息队列、微服务和高并发基础。",
    10: "建立持续评测、发布、观测和线上故障定位体系。",
    11: "把知识映射到上线、运维、安全和代码 Agent 实践。",
    12: "把真实报价中台压缩为可复述的面试答案与模拟题。",
    13: "补齐 Linux、容器、Kubernetes、CI/CD 与可观测性等云原生工程能力。",
    14: "补齐算法、操作系统、网络、Python、并发、数据库原理与现场编码训练。",
}


def add_cover(document: Document, total_notes: int) -> None:
    top = document.add_paragraph()
    top.paragraph_format.space_after = Pt(24)
    label = top.add_run("PRODUCTION AI AGENT FIELD GUIDE")
    set_run_font(label, size=9, bold=True, color=BLUE)
    label.font.all_caps = True

    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(5)
    run = title.add_run("AI Agent 工程师 / FDE")
    set_run_font(run, size=27, bold=True, color=DARK_BLUE)

    title2 = document.add_paragraph()
    title2.paragraph_format.space_after = Pt(17)
    run = title2.add_run("综合学习与模拟面试手册")
    set_run_font(run, size=23, bold=True, color=BLUE)

    rule = document.add_paragraph()
    rule.paragraph_format.space_after = Pt(16)
    set_paragraph_border(rule, "bottom", BLUE, size=14, space=1)

    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(21)
    run = subtitle.add_run(
        "Agent · RAG · 模型应用 · 后端 · 云原生 · 评测安全治理 · 报价中台实战"
    )
    set_run_font(run, size=12, color=TEXT)

    table = document.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [2500, 6860])
    metadata = [
        ("用途", "系统复习、面试表达、项目复盘与查漏补缺"),
        ("收录", f"{total_notes} 篇已整理 Markdown 笔记"),
        ("重点", "生产级 Agent 工程 + AI 智能报价中台真实映射"),
        ("版本", "2026.07 · 完整版（含三大补充模块与第一轮模拟面试）"),
    ]
    for row_index, (label_text, value) in enumerate(metadata):
        left, right = table.rows[row_index].cells
        set_cell_shading(left, LIGHT_BLUE)
        set_cell_shading(right, WHITE if row_index % 2 == 0 else "FAFBFC")
        left.text = ""
        right.text = ""
        left_p = left.paragraphs[0]
        right_p = right.paragraphs[0]
        left_p.paragraph_format.space_after = Pt(0)
        right_p.paragraph_format.space_after = Pt(0)
        left_run = left_p.add_run(label_text)
        right_run = right_p.add_run(value)
        set_run_font(left_run, size=9.5, bold=True, color=DARK_BLUE)
        set_run_font(right_run, size=9.5, color=TEXT)

    spacer = document.add_paragraph()
    spacer.paragraph_format.space_before = Pt(27)
    spacer.paragraph_format.space_after = Pt(4)
    run = spacer.add_run("学习原则")
    set_run_font(run, size=10, bold=True, color=MUTED)
    principle = document.add_paragraph()
    principle.paragraph_format.space_after = Pt(0)
    run = principle.add_run("先讲结论，再讲方案；用真实代码、指标、故障与边界证明能力。")
    set_run_font(run, size=12, bold=True, color=DARK_BLUE)
    principle.add_run().add_break(WD_BREAK.PAGE)


def add_intro_and_toc(document: Document, chapters: list[Chapter], bookmark_id: int) -> int:
    heading = document.add_paragraph(style="Heading 1")
    heading.add_run("使用说明与目录")
    add_bookmark(heading, "contents", bookmark_id)
    bookmark_id += 1

    paragraph = document.add_paragraph()
    paragraph.add_run(
        "本手册按学习索引合并全部笔记，删除 YAML 元数据，但保留正文层级、表格、代码块与关联链接。"
        "目录可点击跳转，Word 的“导航窗格”也可按标题快速检索。"
    )

    heading = document.add_paragraph(style="Heading 2")
    heading.add_run("建议复习顺序")
    for text in (
        "第一遍：01 → 14 → 04 → 03 → 08 → 09 → 13 → 10 → 12，先形成端到端主线。",
        "第二遍：结合报价中台，重点复述模型推理、异步任务、RAG 证据链、Agent 编排、云原生运维与安全治理。",
        "第三遍：完成第 12 章第一轮模拟面试，按评分表定位薄弱项。",
        "面试前：只看每章结论、项目映射、故障案例、指标与取舍边界。",
    ):
        paragraph = document.add_paragraph()
        apply_numbering(paragraph, add_numbering_definition(document, False), 0)
        add_inline_content(paragraph, text, None, {})

    heading = document.add_paragraph(style="Heading 2")
    heading.add_run("全书目录（点击标题跳转）")
    for chapter in chapters:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(8)
        paragraph.paragraph_format.space_after = Pt(2)
        add_internal_hyperlink(
            paragraph,
            f"{chapter.index:02d}. {chapter.title}（{len(chapter.notes)} 篇）",
            chapter.bookmark,
            DARK_BLUE,
        )
        for note in chapter.notes:
            note_p = document.add_paragraph()
            note_p.paragraph_format.left_indent = Inches(0.2)
            note_p.paragraph_format.first_line_indent = Inches(-0.12)
            note_p.paragraph_format.space_after = Pt(1.5)
            prefix = "• "
            run = note_p.add_run(prefix)
            set_run_font(run, size=8.5, color=MUTED)
            add_internal_hyperlink(note_p, note.title, note.bookmark)
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    return bookmark_id


def add_note_source_line(document: Document, note: Note) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(7)
    run = paragraph.add_run(f"源笔记：{note.relative_source}")
    set_run_font(run, size=8, color=MUTED, italic=True)


def build_handbook(notes_root: Path, output_path: Path) -> None:
    chapters = parse_readme(notes_root)
    total_notes = sum(len(chapter.notes) for chapter in chapters)
    bookmark_by_source = {
        note.source.resolve(): note.bookmark
        for chapter in chapters
        for note in chapter.notes
    }

    document = Document()
    configure_styles(document)
    configure_page(document)
    configure_headers_and_footers(document)

    properties = document.core_properties
    properties.title = "AI Agent 工程师 / FDE 综合学习与模拟面试手册"
    properties.subject = "生产级 Agent、RAG、模型应用、后端、云原生、评测安全治理与报价中台面试复习"
    properties.author = "Codex"
    properties.keywords = "AI Agent, FDE, RAG, LLMOps, Backend, Interview"
    properties.comments = f"由 {total_notes} 篇 Markdown 学习笔记合并生成"

    settings = document.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    add_cover(document, total_notes)
    bookmark_id = 1
    bookmark_id = add_intro_and_toc(document, chapters, bookmark_id)

    renderer = MarkdownRenderer(document, bookmark_by_source)
    for chapter in chapters:
        chapter_heading = document.add_paragraph(style="Heading 1")
        chapter_heading.paragraph_format.page_break_before = True
        chapter_heading.add_run(f"{chapter.index:02d}. {chapter.title}")
        add_bookmark(chapter_heading, chapter.bookmark, bookmark_id)
        bookmark_id += 1

        description = document.add_paragraph()
        description.paragraph_format.space_after = Pt(12)
        run = description.add_run(CHAPTER_DESCRIPTIONS.get(chapter.index, ""))
        set_run_font(run, size=10.5, color=MUTED, italic=True)

        last_subcategory: str | None = None
        for note in chapter.notes:
            note_level = 2
            if note.subcategory:
                if note.subcategory != last_subcategory:
                    subcategory_heading = document.add_paragraph(style="Heading 2")
                    subcategory_heading.add_run(note.subcategory)
                    last_subcategory = note.subcategory
                note_level = 3

            note_heading = document.add_paragraph(style=f"Heading {note_level}")
            note_heading.add_run(note.title)
            add_bookmark(note_heading, note.bookmark, bookmark_id)
            bookmark_id += 1
            add_note_source_line(document, note)

            source_text = note.source.read_text(encoding="utf-8")
            renderer.render(source_text, note.source, note_level)

            back = document.add_paragraph()
            back.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            back.paragraph_format.space_before = Pt(4)
            back.paragraph_format.space_after = Pt(8)
            add_internal_hyperlink(back, "↑ 返回目录", "contents", MUTED)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)

    reopened = Document(output_path)
    if len(reopened.paragraphs) < 1000:
        raise RuntimeError(f"生成结果段落数异常：{len(reopened.paragraphs)}")
    actual_titles = {
        paragraph.text.strip()
        for paragraph in reopened.paragraphs
        if paragraph.style and paragraph.style.name.startswith("Heading")
    }
    missing = [note.title for chapter in chapters for note in chapter.notes if note.title not in actual_titles]
    if missing:
        raise RuntimeError(f"生成结果缺少 {len(missing)} 篇笔记标题：{missing[:5]}")

    print(f"output={output_path.resolve()}")
    print(f"chapters={len(chapters)}")
    print(f"notes={total_notes}")
    print(f"paragraphs={len(reopened.paragraphs)}")
    print(f"tables={len(reopened.tables)}")
    print(f"generated_at={datetime.now().isoformat(timespec='seconds')}")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Build the consolidated AI Agent learning handbook.")
    parser.add_argument("--notes-root", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    return parser.parse_args()


if __name__ == "__main__":
    args = parse_args()
    build_handbook(args.notes_root.resolve(), args.output.resolve())
